from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2C3E50')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    return p

def set_font(run, name='Garamond', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# --- HEADER ---
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_para.paragraph_format.space_after = Pt(2)
r = name_para.add_run("Victoria Bernadskiy")
set_font(r, size=22, bold=True, color=(30, 30, 30))

contact_para = doc.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_para.paragraph_format.space_after = Pt(8)
r = contact_para.add_run("Staten Island, NY  |  (718) 200-5932  |  vbernadskiy0531@gmail.com")
set_font(r, size=10, color=(80, 80, 80))

add_horizontal_rule(doc)

def section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    set_font(r, size=11, bold=True, color=(44, 62, 80))
    add_horizontal_rule(doc)

def subheading_line(doc, left, right=None, left_italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(left)
    set_font(r1, size=11, bold=True, italic=left_italic)
    if right:
        tab_stop = p.paragraph_format
        r2 = p.add_run(f"  —  {right}")
        set_font(r2, size=10, italic=True, color=(100, 100, 100))
    return p

def detail_line(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    set_font(r, size=10, italic=italic, color=(60, 60, 60))
    return p

def bullet_item(doc, text, indent_level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.2 * indent_level)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    set_font(r, size=10)
    return p

# --- EDUCATION ---
section_heading(doc, "Education")

subheading_line(doc, "Macaulay Honors College at Hunter College", "New York, NY")
detail_line(doc, "Bachelor of Arts in Economics and Environmental Science  |  In Progress  |  Expected May 2028")
bullet_item(doc, "GPA: 3.95 / 4.0  —  Macaulay Honors Scholar, National Honor Society")
bullet_item(doc, "Focus: Environmental Economics, Sustainability Policy, Quantitative Analysis, and Geospatial Systems")
bullet_item(doc, "Relevant Coursework: Principles of Micro/Macroeconomics, Environmental Economics, Statistics, "
                 "Weather and Climate, Geographic Information Systems (GIS), Computer Aided Design, Audio Visual Engineering")

# --- TECHNICAL & ACADEMIC PROJECTS ---
section_heading(doc, "Technical & Academic Projects")

subheading_line(doc, "Campus Sustainability Guide — Hunter Sustainability Council", "Sep 2025 – Present")
bullet_item(doc, "Researching and authoring a comprehensive campus guide outlining actionable steps for "
                 "Hunter College Waste Management, Catering, and student clubs to reduce climate impact.")
bullet_item(doc, "Analyzed sustainability initiatives at 20+ colleges nationwide to benchmark best practices "
                 "and formulate evidence-based recommendations.")

subheading_line(doc, "9STARS Energy Database & Strategic Planning Analysis — The New School Intern", "Mar 2025 – May 2025")
bullet_item(doc, "Gathered and entered energy-usage data for 16 campus buildings into the 9STARS Sustainability "
                 "Tracking, Assessment & Rating System institutional database.")
bullet_item(doc, "Analyzed energy-consumption datasets using Excel to identify usage trends, anomalies, and "
                 "correlations supporting long-range infrastructure planning under NYC Local Laws 84, 88, and 97.")
bullet_item(doc, "Created an interactive sustainability resource map and a covered-buildings compliance list "
                 "for internal use across departments.")

# --- PROFESSIONAL EXPERIENCE ---
section_heading(doc, "Professional Experience")

subheading_line(doc, "The New School", "New York, NY  |  Mar 2025 – May 2025")
detail_line(doc, "Energy and Sustainability Intern", italic=True)
bullet_item(doc, "Designed and implemented 3 sustainability projects: a student sustainability guide, an "
                 "interactive resource map, and a Local Law compliance tracker.")
bullet_item(doc, "Benchmarked sustainability efforts across 12 universities to establish campus standards "
                 "across 8 categories for a new institutional guide.")
bullet_item(doc, "Compiled a shortlist of 6 firms (from 20 researched) to conduct climate resiliency assessments, "
                 "supporting procurement decisions.")

subheading_line(doc, "STEM In Action", "New York, NY  |  Sep 2025 – Present")
detail_line(doc, "After-School STEM Instructor", italic=True)
bullet_item(doc, "Design and deliver 9 hands-on STEM modules per semester for elementary students ages 5–11, "
                 "using 18 original PowerPoint presentations and projects to introduce core engineering concepts.")
bullet_item(doc, "Manage classrooms of 12–20 students, facilitating collaborative learning and adapting "
                 "lessons to diverse age groups and skill levels.")

subheading_line(doc, "United Activities Unlimited", "Staten Island, NY  |  Jun 2025 – Aug 2025")
detail_line(doc, "Day Camp Counselor", italic=True)
bullet_item(doc, "Organized and led 20+ weekly recreational, arts & crafts, and team-building activities "
                 "for children ages 12–14.")
bullet_item(doc, "Supervised groups of 40–60 campers daily, ensuring safety, inclusion, and active participation.")

# --- EXTRACURRICULAR ACHIEVEMENTS ---
section_heading(doc, "Extracurricular Achievements")

subheading_line(doc, "Summer Design Institute — Videographer / Team Member", "Jul 2023 – Aug 2024")
bullet_item(doc, "Operated drones and handheld cameras to capture 6+ hours of broadcast-quality footage and "
                 "2 hours of B-roll for 8 morning-show segments.")
bullet_item(doc, "Performed post-production editing in Adobe Premiere Pro and After Effects for 4 segments "
                 "across school news, history, and community features.")
bullet_item(doc, "Managed A/V equipment inventory of 35+ cameras, microphones, lighting rigs, soundboards, "
                 "and teleprompters; collaborated with a 7-person production team.")

# --- SKILLS ---
section_heading(doc, "Skills")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Technical:  ")
set_font(r, bold=True, size=10)
r2 = p.add_run("MS Excel, PowerPoint, HTML, JavaScript, Adobe Premiere Pro, Adobe Photoshop, AutoCAD")
set_font(r2, size=10)

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after = Pt(2)
r = p2.add_run("Certifications:  ")
set_font(r, bold=True, size=10)
r2 = p2.add_run("Computer Aided Design (Sep 2022)")
set_font(r2, size=10)

p3 = doc.add_paragraph()
p3.paragraph_format.space_before = Pt(2)
p3.paragraph_format.space_after = Pt(2)
r = p3.add_run("Languages:  ")
set_font(r, bold=True, size=10)
r2 = p3.add_run("English (Native), Russian (Fluent)")
set_font(r2, size=10)

# --- HONORS & LEADERSHIP ---
section_heading(doc, "Honors & Leadership")

bullet_item(doc, "Macaulay Honors Scholar — Awarded to top-performing students at CUNY Macaulay Honors College")
bullet_item(doc, "National Honor Society Member")
bullet_item(doc, "Academic Honor Roll — All Semesters (GPA 3.95)")
bullet_item(doc, "Hunter Sustainability Council — Active Member & Research Contributor (Sep 2025–Present)")

# Save
output_path = os.path.expanduser("~/Desktop/Victoria_Bernadskiy_Resume.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
