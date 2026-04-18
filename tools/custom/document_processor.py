"""
Document Processor — create and reformat documents in DOCX, PDF, TXT, and MD formats.

Gemma4-compatible: all formatting is handled in Python; the LLM only provides content.
"""
import os
import re
from pathlib import Path
from datetime import datetime

_OUTPUT_DIR = Path.home() / "Documents" / "HUBERT"


def _ensure_output_dir():
    _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _resolve_output_path(filename: str, fmt: str) -> Path:
    """Build the output path, appending .{fmt} if not already present."""
    _ensure_output_dir()
    stem = Path(filename).stem or filename
    return _OUTPUT_DIR / f"{stem}.{fmt}"


# ── DOCX writer ──────────────────────────────────────────────────────────────

def _write_docx(content: str, out_path: Path, title: str = "") -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Optional title
    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse content: lines starting with # / ## / ### become headings
    for line in content.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 60)
        elif stripped == "":
            doc.add_paragraph("")
        else:
            # Bold (**text**) and italic (*text*) inline rendering
            p = doc.add_paragraph()
            _render_inline(p, stripped)

    doc.save(str(out_path))
    return out_path


def _render_inline(paragraph, text: str):
    """Render **bold** and *italic* inline markdown into a docx paragraph."""
    from docx.shared import Pt
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ── PDF writer ───────────────────────────────────────────────────────────────

def _write_pdf(content: str, out_path: Path, title: str = "") -> Path:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    if title:
        pdf.set_font("Helvetica", "B", 18)
        pdf.multi_cell(0, 10, title, align="C")
        pdf.ln(6)

    for line in content.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, stripped[4:])
            pdf.ln(2)
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 15)
            pdf.multi_cell(0, 9, stripped[3:])
            pdf.ln(3)
        elif stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 17)
            pdf.multi_cell(0, 10, stripped[2:])
            pdf.ln(4)
        elif stripped.startswith("---"):
            pdf.set_draw_color(180, 180, 180)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(4)
        elif stripped == "":
            pdf.ln(4)
        else:
            # Strip markdown bold/italic markers for PDF plain text
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            clean = re.sub(r'\*(.*?)\*', r'\1', clean)
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, clean)

    pdf.output(str(out_path))
    return out_path


# ── Tool handlers ─────────────────────────────────────────────────────────────

def _open_file(path: str):
    """Open a file in the default OS application (non-blocking)."""
    import subprocess, platform
    try:
        if platform.system() == "Darwin":
            subprocess.Popen(["open", path])
        elif platform.system() == "Windows":
            subprocess.Popen(["start", "", path], shell=True)
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception:
        pass


def _create_document(params: dict) -> str:
    content  = params.get("content", "").strip()
    filename = params.get("filename", f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    fmt      = params.get("format", "docx").lower().lstrip(".")
    title    = params.get("title", "")
    auto_open = params.get("auto_open", True)

    if not content:
        return "Error: no content provided."
    if fmt not in ("docx", "pdf", "txt", "md"):
        return f"Error: unsupported format '{fmt}'. Use docx, pdf, txt, or md."

    out_path = _resolve_output_path(filename, fmt)

    try:
        if fmt == "docx":
            _write_docx(content, out_path, title=title)
        elif fmt == "pdf":
            _write_pdf(content, out_path, title=title)
        else:
            out_path.write_text(content, encoding="utf-8")

        if auto_open:
            _open_file(str(out_path))
        return f"Document saved: {out_path}"
    except Exception as e:
        return f"Error creating document: {e}"


def _reformat_document(params: dict) -> str:
    """
    Read an existing document and write it in a new format.
    No LLM involved — pure format conversion.
    """
    source_path = params.get("source_path", "").strip()
    target_fmt  = params.get("format", "docx").lower().lstrip(".")
    filename    = params.get("filename", "")
    title       = params.get("title", "")

    if not source_path or not Path(source_path).exists():
        return f"Error: source file not found: {source_path}"

    # Extract text from source
    from file_upload_utils import extract_text
    content = extract_text(source_path, max_chars=50_000)
    if not content:
        return f"Error: could not extract text from {source_path}"

    out_name = filename or Path(source_path).stem
    out_path = _resolve_output_path(out_name, target_fmt)

    try:
        if target_fmt == "docx":
            _write_docx(content, out_path, title=title or Path(source_path).stem)
        elif target_fmt == "pdf":
            _write_pdf(content, out_path, title=title or Path(source_path).stem)
        else:
            out_path.write_text(content, encoding="utf-8")

        _open_file(str(out_path))
        return f"Reformatted document saved: {out_path}"
    except Exception as e:
        return f"Error reformatting document: {e}"


def _combine_documents(params: dict) -> str:
    """
    Combine multiple documents into one. Pass their extracted text blocks as a list.
    """
    sources  = params.get("source_paths", [])
    fmt      = params.get("format", "docx").lower().lstrip(".")
    filename = params.get("filename", f"combined_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    title    = params.get("title", "Combined Document")

    if not sources:
        return "Error: no source_paths provided."

    from file_upload_utils import extract_text
    parts = []
    for sp in sources:
        if not Path(sp).exists():
            parts.append(f"# {Path(sp).name}\n[File not found]")
            continue
        text = extract_text(sp, max_chars=20_000)
        parts.append(f"# {Path(sp).name}\n\n{text or '[Could not extract text]'}")

    combined = "\n\n---\n\n".join(parts)
    out_path = _resolve_output_path(filename, fmt)

    try:
        if fmt == "docx":
            _write_docx(combined, out_path, title=title)
        elif fmt == "pdf":
            _write_pdf(combined, out_path, title=title)
        else:
            out_path.write_text(combined, encoding="utf-8")

        _open_file(str(out_path))
        return f"Combined document saved: {out_path} ({len(sources)} sources)"
    except Exception as e:
        return f"Error combining documents: {e}"


def _open_google_doc(params: dict) -> str:
    """
    Open Google Docs in the browser and copy content to clipboard so the user
    can paste it into a new document. Optionally uploads a local file path.
    """
    import webbrowser
    content   = params.get("content", "")
    file_path = params.get("file_path", "")

    try:
        import pyperclip
        if content:
            pyperclip.copy(content)
        elif file_path and Path(file_path).exists():
            from file_upload_utils import extract_text
            text = extract_text(file_path)
            if text:
                pyperclip.copy(text)
        clipboard_msg = " Content copied to clipboard — paste with Cmd+V (Mac) or Ctrl+V (Windows)." if content or file_path else ""
    except Exception:
        clipboard_msg = ""

    webbrowser.open("https://docs.google.com/document/create")
    return f"Opened Google Docs in browser.{clipboard_msg}"


# ── Tool registration ─────────────────────────────────────────────────────────

TOOLS = [
    (
        {
            "name": "create_document",
            "description": (
                "Create a new document file (docx, pdf, txt, md) from text content. "
                "Supports markdown headings (# / ## / ###), bold (**text**), italic (*text*). "
                "Saves to ~/Documents/HUBERT/. Returns the saved file path."
            ),
            "input_schema": {
                "type": "object",
                "properties": {
                    "content":  {"type": "string",  "description": "Full text/markdown content of the document"},
                    "filename": {"type": "string",  "description": "Output filename without extension (e.g. 'my_report')"},
                    "format":   {"type": "string",  "enum": ["docx", "pdf", "txt", "md"],
                                 "description": "Output format"},
                    "title":    {"type": "string",  "description": "Optional document title shown at the top"},
                },
                "required": ["content", "filename", "format"],
            },
        },
        _create_document,
    ),
    (
        {
            "name": "reformat_document",
            "description": (
                "Convert an existing document to a different format (docx, pdf, txt, md). "
                "Reads the source file, extracts text, and writes it in the target format."
            ),
            "input_schema": {
                "type": "object",
                "properties": {
                    "source_path": {"type": "string", "description": "Full path to the source document"},
                    "format":      {"type": "string", "enum": ["docx", "pdf", "txt", "md"],
                                   "description": "Target format"},
                    "filename":    {"type": "string", "description": "Output filename without extension (optional)"},
                    "title":       {"type": "string", "description": "Optional title for the output document"},
                },
                "required": ["source_path", "format"],
            },
        },
        _reformat_document,
    ),
    (
        {
            "name": "open_google_doc",
            "description": (
                "Open a new Google Doc in the browser. "
                "If content or a file_path is provided, it is copied to the clipboard "
                "so the user can paste it into the new document."
            ),
            "input_schema": {
                "type": "object",
                "properties": {
                    "content":   {"type": "string",
                                  "description": "Text content to copy to clipboard for pasting"},
                    "file_path": {"type": "string",
                                  "description": "Path to a local file whose text will be copied to clipboard"},
                },
                "required": [],
            },
        },
        _open_google_doc,
    ),
    (
        {
            "name": "combine_documents",
            "description": (
                "Merge multiple documents into one file. "
                "Extracts text from each source and combines them with section headers."
            ),
            "input_schema": {
                "type": "object",
                "properties": {
                    "source_paths": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "List of full paths to source documents",
                    },
                    "format":   {"type": "string", "enum": ["docx", "pdf", "txt", "md"],
                                 "description": "Output format"},
                    "filename": {"type": "string", "description": "Output filename without extension"},
                    "title":    {"type": "string", "description": "Title for the combined document"},
                },
                "required": ["source_paths", "format"],
            },
        },
        _combine_documents,
    ),
]
